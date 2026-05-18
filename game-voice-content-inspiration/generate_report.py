import sys
import subprocess
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='Microsoft YaHei', font_size=11, bold=False):
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def markdown_to_docx(md_file, docx_file):
    doc = Document()
    
    # Set default font for document
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip('\n')
        if not line.strip():
            continue
            
        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:]
            p = doc.add_heading(level=1)
            run = p.add_run(text)
            set_chinese_font(run, font_size=18, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Heading 2
        elif line.startswith('## '):
            text = line[3:]
            p = doc.add_heading(level=2)
            run = p.add_run(text)
            set_chinese_font(run, font_size=16, bold=True)
        # Heading 3
        elif line.startswith('### '):
            text = line[4:]
            p = doc.add_heading(level=3)
            run = p.add_run(text)
            set_chinese_font(run, font_size=14, bold=True)
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            set_chinese_font(run, font_size=11)
        # Numbered list
        elif line.strip() and line.strip()[0].isdigit() and '. ' in line.strip()[:4]:
            text = line.strip()
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(text)
            set_chinese_font(run, font_size=11)
        # Table separator or normal text
        elif line.strip().startswith('|'):
            # Skip table separators
            if '---' in line:
                continue
            text = line.strip().strip('|').replace('|', ' | ')
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_chinese_font(run, font_size=10)
        # Bold text handling
        else:
            p = doc.add_paragraph()
            # Parse inline bold **text**
            parts = line.split('**')
            for i, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    set_chinese_font(run, font_size=11, bold=(i % 2 == 1))
    
    doc.save(docx_file)
    print(f"Document saved to: {docx_file}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python generate_report.py <input.md> <output.docx>")
        sys.exit(1)
    
    markdown_to_docx(sys.argv[1], sys.argv[2])
