#!/usr/bin/env python3
"""
报告导出为 Word 文档
"""

import sys
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("[错误] 未安装 python-docx，请先运行: pip install python-docx")
    sys.exit(1)

sys.path.insert(0, str(Path(__file__).parent))
from config_manager import REPORTS_DIR, ensure_dirs


def get_desktop_path():
    """获取用户桌面路径"""
    return Path.home() / "Desktop"


def markdown_to_word(md_path, output_path=None):
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 标题
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
            i += 1
            continue

        # 引用块 / 摘要信息
        if line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # 分隔线
        if line.strip() == "---":
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # 表格
        if line.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            if len(rows) >= 2:
                # 解析表头
                header = [c.strip() for c in rows[0].split("|") if c.strip()]
                table = doc.add_table(rows=1, cols=len(header))
                table.style = 'Table Grid'
                for j, h in enumerate(header):
                    table.rows[0].cells[j].text = h
                # 跳过表格分隔行，解析数据行
                for r in rows[2:]:
                    cells = [c.strip() for c in r.split("|") if c.strip()]
                    if cells:
                        row_cells = table.add_row().cells
                        for j, c in enumerate(cells[:len(header)]):
                            row_cells[j].text = c
            continue

        # 列表项
        if line.strip().startswith("- "):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s", line):
            p = doc.add_paragraph(re.sub(r"^\s*\d+\.\s", "", line), style='List Number')
            i += 1
            continue

        # 粗体段落（如 **提及该话题的KOL**:）
        m = re.match(r"\*\*(.*?)\*\*", line)
        if m and line.strip().endswith(":"):
            doc.add_paragraph(line.strip(), style='Normal')
            i += 1
            continue

        # 普通文本
        if line.strip():
            doc.add_paragraph(line.strip())

        i += 1

    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    if not output_path:
        timestamp = datetime.now().strftime("%Y-%m-%d")
        output_path = REPORTS_DIR / f"{timestamp}_report.docx"

    ensure_dirs()
    doc.save(output_path)
    print(f"[成功] Word报告已导出: {output_path}")
    return output_path


def main():
    args = sys.argv[1:]
    md_file = None
    output_file = None

    if "--input" in args:
        idx = args.index("--input")
        md_file = Path(args[idx + 1]) if idx + 1 < len(args) else None

    if "--output" in args:
        idx = args.index("--output")
        output_file = Path(args[idx + 1]) if idx + 1 < len(args) else None

    if not md_file:
        md_files = sorted(REPORTS_DIR.glob("*_report_v2.md"), reverse=True)
        if not md_files:
            print("[错误] 未找到报告文件")
            sys.exit(1)
        md_file = md_files[0]

    markdown_to_word(md_file, output_file)


if __name__ == "__main__":
    main()
